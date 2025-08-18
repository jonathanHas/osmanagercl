#!/usr/bin/env python3
"""
Parse .doc file and extract text content
"""

import sys
import os
from docx import Document
import subprocess
import tempfile

def parse_doc_file(file_path):
    """Parse a .doc or .docx file and return its text content"""
    
    if not os.path.exists(file_path):
        print(f"Error: File '{file_path}' not found")
        return None
    
    file_ext = os.path.splitext(file_path)[1].lower()
    
    if file_ext == '.docx':
        # Handle .docx files directly
        try:
            doc = Document(file_path)
            text = []
            for paragraph in doc.paragraphs:
                text.append(paragraph.text)
            return '\n'.join(text)
        except Exception as e:
            print(f"Error reading .docx file: {e}")
            return None
    
    elif file_ext == '.doc':
        # Convert .doc to .docx using LibreOffice, then parse
        try:
            # Create temporary directory for conversion
            with tempfile.TemporaryDirectory() as temp_dir:
                # Convert .doc to .docx using LibreOffice
                cmd = [
                    'libreoffice', '--headless', '--convert-to', 'docx',
                    '--outdir', temp_dir, file_path
                ]
                
                result = subprocess.run(cmd, capture_output=True, text=True)
                
                if result.returncode != 0:
                    print(f"Error converting file: {result.stderr}")
                    print("Make sure LibreOffice is installed (sudo apt-get install libreoffice)")
                    return None
                
                # Find the converted file
                base_name = os.path.splitext(os.path.basename(file_path))[0]
                docx_path = os.path.join(temp_dir, f"{base_name}.docx")
                
                if not os.path.exists(docx_path):
                    print("Error: Converted file not found")
                    return None
                
                # Parse the converted .docx file
                doc = Document(docx_path)
                text = []
                for paragraph in doc.paragraphs:
                    text.append(paragraph.text)
                return '\n'.join(text)
                
        except Exception as e:
            print(f"Error processing .doc file: {e}")
            return None
    
    else:
        print(f"Error: Unsupported file format '{file_ext}'")
        return None

def main():
    if len(sys.argv) < 2:
        print("Usage: python parse_doc_file.py <path_to_doc_file>")
        sys.exit(1)
    
    file_path = sys.argv[1]
    content = parse_doc_file(file_path)
    
    if content:
        print(content)

if __name__ == "__main__":
    main()